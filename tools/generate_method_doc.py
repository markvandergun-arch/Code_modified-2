from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "methode_energieplanner.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)


def set_run(run, *, size: float | None = None, bold: bool | None = None, color=None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str, *, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_small_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    hdr[0].text = "Onderdeel"
    hdr[1].text = "Rol in het model"
    for label, description in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = description
    for row in table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(12.0)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(9)


def add_three_col_table(doc: Document, headers: tuple[str, str, str], rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    widths = [Cm(4.0), Cm(7.0), Cm(5.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(8.5)


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK_BLUE


def add_model_section(
    doc: Document,
    title: str,
    *,
    purpose: str,
    code_basis: str,
    method: list[str],
    formulas: list[str],
    outputs: list[str],
    motivation: str,
) -> None:
    doc.add_heading(title, level=2)
    add_paragraph(doc, purpose)
    add_three_col_table(
        doc,
        ("Aspect", "Beschrijving", "Belangrijkste code-uitvoer"),
        [
            ("Aansluiting op code", code_basis, ", ".join(outputs)),
            ("Motivatie", motivation, ""),
        ],
    )
    add_bullets(doc, method)
    for formula in formulas:
        add_formula(doc, formula)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_doc() -> Document:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Methode en onderbouwing energieplanner")
    set_run(run, size=22, bold=True, color=DARK_BLUE)
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Werkdocument voor stapsgewijze uitwerking van de modelmethode").italic = True
    subtitle.runs[0].font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    add_paragraph(
        doc,
        "Dit document wordt incrementeel opgebouwd. Eerst wordt het doel van het model vastgelegd. "
        "Daarna worden per deelmodel de aannames, rekenstappen, interpretatie en bronnen toegevoegd. "
        "De inhoud is bedoeld als basis voor een later tabblad 'Methode & uitleg' in de app en eventueel "
        "voor een downloadbare Word- of PDF-handleiding.",
    )

    doc.add_heading("Implementatieplan voor het methodedocument", level=1)
    add_paragraph(
        doc,
        "De methode-uitleg wordt bewust in stappen opgebouwd. Zo kan per onderdeel inhoudelijk worden gecontroleerd "
        "of de tekst klopt met de code, begrijpelijk is voor consultants en voldoende onderbouwd is voor klantgesprekken."
    )
    add_numbered(
        doc,
        [
            "Doel van het model: vastleggen waarvoor de app bedoeld is, welke energiestromen worden onderscheiden en welke beslissingen de app ondersteunt.",
            "Modeloverzicht: beschrijven hoe gebouwvraag, verbruik, opwek, warmtebronnen, opslag en netbelasting elkaar opvolgen in de simulatie.",
            "Deelmodellen: per onderdeel uitleggen wat wordt berekend, welke invoer nodig is, welke aannames worden gemaakt en hoe resultaten moeten worden gelezen.",
            "Modelchecks en beperkingen: uitleggen wat de kwaliteitschecks betekenen, wat buiten scope valt en hoe onzekerheden moeten worden geinterpreteerd.",
            "Bronnen en onderbouwing: per thema bronnen toevoegen, bijvoorbeeld voor gebouwenergie, PV, mobiliteit, warmtepompen, opslag en netcapaciteitsindicatoren.",
            "Vertaling naar de app: de definitieve tekst verwerken in een tabblad 'Methode & uitleg' met expanders per onderdeel en eventueel een downloadbaar methoderapport.",
        ],
    )

    doc.add_heading("Voorgestelde documentstructuur", level=1)
    add_small_table(
        doc,
        [
            ("1. Doel van het model", "Waarom de app bestaat en welke vraag hij beantwoordt."),
            ("2. Afbakening en toepassingsgebied", "Wat het model wel en niet pretendeert te doen."),
            ("3. Rekenstructuur", "Volgorde van energiestromen: vraag, opwek, warmteketen, opslag, net."),
            ("4. Deelmodellen", "Gebouw, elektrisch verbruik, processen, mobiliteit, PV, WKK, warmte, opslag en validatie."),
            ("5. Resultaatindicatoren", "Piek netimport, load factor, zelfvoorziening, ongedekte warmte, teruglevering en onbenutte warmte."),
            ("6. Modelchecks", "Warmtebalans, negatieve fysica, vermogensgrenzen en validatie met meetdata."),
            ("7. Aannames en beperkingen", "Geen kostenoptimalisatie, geen hydraulisch/detail-installatiemodel, afhankelijkheid van invoer en weerdata."),
            ("8. Bronnen", "Literatuur en technische documentatie per modelonderdeel."),
        ],
    )

    doc.add_heading("1. Doel van het model", level=1)
    add_paragraph(
        doc,
        "De energieplanner heeft als doel om de energiestromen van een individueel gebouw inzichtelijk te maken. "
        "Het model splitst de energievraag en energielevering expliciet op in drie dragers: thermische energie, "
        "elektrische energie en brandstofenergie. Door deze dragers gescheiden te modelleren kan de gebruiker zien "
        "welke vraag uit het gebouw zelf komt, welke vraag elektrisch wordt ingevuld, welke warmtebronnen worden gebruikt "
        "en welk deel uiteindelijk nog via het elektriciteitsnet of via brandstof wordt geleverd."
    )
    add_paragraph(
        doc,
        "De app is daarmee primair bedoeld als scenariotool voor consultants. Een gebruiker kan instellingen aanpassen "
        "voor gebouwgebruik, elektrische lasten, processen, mobiliteit, zonnepanelen, WKK, warmtepompen, ketels, warmtenet, "
        "batterijen en warmteopslag. Vervolgens laat de simulatie zien hoe deze keuzes doorwerken in de totale energiebalans, "
        "de netbelasting, de warmtevoorziening, teruglevering, opslaggebruik en eventuele ongedekte warmte."
    )

    doc.add_heading("1.1 Centrale ontwerpvraag", level=2)
    add_paragraph(
        doc,
        "De centrale ontwerpvraag is niet alleen hoeveel energie een gebouw per jaar gebruikt, maar vooral wanneer die energie "
        "nodig is en via welke drager die geleverd wordt. Dat onderscheid is essentieel voor twee praktische vraagstukken:"
    )
    add_bullets(
        doc,
        [
            "Kan een gebouw geheel of gedeeltelijk van het gas af zonder dat de warmtevraag ongedekt blijft?",
            "Kan een gebouw functioneren binnen een beperkte netaansluiting, eventueel met lokale opwek, opslag en slim laden?",
        ],
    )
    add_paragraph(
        doc,
        "Daarom rekent het model niet alleen met jaarvolumes, maar met tijdreeksen. Pieken, gelijktijdigheid en seizoenspatronen "
        "zijn minstens zo belangrijk als het jaarlijkse totaal. Een gebouw kan op jaarbasis voldoende lokale opwek hebben, maar "
        "toch een netprobleem houden als vraag en opwek op andere momenten optreden."
    )

    doc.add_heading("1.2 Energiedragers in het model", level=2)
    add_small_table(
        doc,
        [
            (
                "Thermische energie",
                "Warmtevraag en koelvraag worden als thermische vraag berekend. Warmtebronnen zoals warmtepomp, WKK-warmte, ketel, warmtenet, referentieverwarming en warmteopslag vullen deze vraag in.",
            ),
            (
                "Elektrische energie",
                "Elektrisch gebouwverbruik, processen, mobiliteit, koeling, warmtepompen, PV, WKK-elektriciteit, batterij en netimport bepalen samen de elektrische balans.",
            ),
            (
                "Brandstofenergie",
                "Brandstofstromen ontstaan vooral bij ketel en WKK. Deze worden apart bijgehouden om gasloosheid en resterend brandstofgebruik zichtbaar te maken.",
            ),
        ],
    )

    doc.add_heading("1.3 Hoe dit aansluit op de code", level=2)
    add_paragraph(
        doc,
        "In de code is deze opzet terug te zien in de totale simulatie. Eerst worden de gebouwgebonden vraag en elektrische lasten "
        "opgebouwd. Daarna worden PV en WKK toegevoegd. Vervolgens wordt de warmtevraag via WKK-warmte, warmteopslag, warmtepomp, "
        "ketel, warmtenet en eventueel referentie-elektrische verwarming afgehandeld. Tot slot wordt de elektrische balans bepaald "
        "met batterijgedrag en netimport/netexport."
    )
    add_bullets(
        doc,
        [
            "Gebouw en lasten: de verbruikscomponenten leveren kolommen zoals warmtevraag, koelvraag, elektrisch basisverbruik, processen, mobiliteit en overig verbruik.",
            "Opwek: zonnepanelen en WKK leveren elektrische opwek; WKK kan daarnaast warmte leveren of onbenutte warmte veroorzaken.",
            "Warmteketen: resterende warmtevraag wordt stap voor stap door opslag, warmtepomp, ketel, warmtenet en referentievoorziening ingevuld.",
            "Opslag en net: batterijgedrag wordt toegepast op de resterende elektrische balans; daarna ontstaan netimport, teruglevering en contractoverschrijding.",
            "Resultaten: KPI's en grafieken tonen piek netimport, jaarverbruik, herkomst elektriciteit, warmtebalans, opslagstromen, gas-/brandstofinput en modelchecks.",
        ],
    )

    doc.add_heading("1.4 Bedoelde toepassing", level=2)
    add_paragraph(
        doc,
        "Het model is bedoeld voor verkennende en vergelijkende analyses. De gebruiker kan scenario's naast elkaar zetten en beoordelen "
        "welke maatregelen bijdragen aan minder gasgebruik, lagere netpieken, betere benutting van lokale opwek of minder ongedekte warmte. "
        "Voorbeelden zijn: meer of anders georienteerde zonnepanelen, een grotere warmtepomp, wel of geen WKK, batterijcapaciteit, warmteopslag "
        "of slim laden van elektrische voertuigen."
    )
    add_paragraph(
        doc,
        "De uitkomst moet worden gelezen als technisch-energetische ondersteuning bij besluitvorming. De app vervangt geen detailontwerp "
        "van installaties, geen kostenoptimalisatie en geen definitieve netstudie, maar helpt wel om vroeg in het proces te zien waar de "
        "belangrijkste knelpunten en oplossingsrichtingen zitten."
    )

    doc.add_heading("1.5 Eerste afbakening", level=2)
    add_bullets(
        doc,
        [
            "Schaalniveau: het model kijkt naar een gebouw of gebouwlocatie, niet naar een volledig gebiedsenergiesysteem.",
            "Tijdsdimensie: de app werkt met tijdreeksen, zodat pieken en seizoenseffecten zichtbaar worden.",
            "Optimalisatie: de huidige logica simuleert gekozen instellingen en dispatchregels; het model kiest niet automatisch de economisch optimale configuratie.",
            "Installatiedetail: de app gebruikt vereenvoudigde technische modellen en is niet bedoeld als hydraulisch, bouwfysisch of elektrotechnisch detailontwerp.",
            "Interpretatie: resultaten zijn sterk afhankelijk van invoerkwaliteit, meetdata, weerdata en gekozen scenario-aannames.",
        ],
    )

    doc.add_heading("2. Modeloverzicht", level=1)
    add_paragraph(
        doc,
        "Het model bestaat uit een vaste rekenketen. De volgorde is belangrijk: eerst wordt bepaald welke vraag het gebouw en "
        "het gebruik veroorzaken, daarna welke lokale opwek en warmtebronnen beschikbaar zijn, vervolgens hoe opslag wordt "
        "ingezet en pas aan het einde wat er van of naar het elektriciteitsnet gaat. Daardoor blijft zichtbaar of een maatregel "
        "de vraag verlaagt, lokale opwek toevoegt, warmte invult, flexibiliteit levert of alleen de resterende netbalans verandert."
    )
    add_three_col_table(
        doc,
        ("Stap", "Rekenblok", "Belangrijkste uitkomst"),
        [
            ("1", "Weerdata en tijdindex", "Simulatie-index met temperatuur, globale zoninstraling en eventueel wind."),
            ("2", "Gebouwmodel", "Thermische warmtevraag Q_heat_kWth en koelvraag Q_cool_kWth."),
            ("3", "Gebruiksmodellen", "Elektrische basislast, processen, mobiliteitsladen en overig verbruik."),
            ("4", "Lokale opwek", "PV-elektriciteit en WKK-elektriciteit plus WKK-warmte."),
            ("5", "Warmteketen", "Invulling van warmtevraag door WKK, warmteopslag, warmtepomp, ketel, warmtenet en referentievoorziening."),
            ("6", "Opslag", "Warmteopslag voor warmteoverschot/tekort en batterij voor elektrische restvraag of overschot."),
            ("7", "Netbalans", "Netimport, teruglevering, contractoverschrijding en stoplichtbeoordeling."),
            ("8", "Validatie en modelchecks", "Controle op warmtebalans, negatieve fysica, capaciteitsgrenzen en vergelijking met meetdata."),
        ],
    )
    add_formula(
        doc,
        "P_load_total = P_base_electric + P_heat_pump_el + P_reference_heat_el"
    )
    add_formula(
        doc,
        "P_residual_before_battery = P_load_total - (P_pv + P_wkk_el)"
    )
    add_formula(
        doc,
        "P_grid_import = max(P_residual_before_battery + P_battery_charge - P_battery_discharge, 0)"
    )
    add_formula(
        doc,
        "P_grid_export = max(-(P_residual_before_battery + P_battery_charge - P_battery_discharge), 0)"
    )
    add_paragraph(
        doc,
        "In woorden: warmte en elektriciteit worden niet op een hoop gegooid. Het gebouw levert eerst thermische vraag; "
        "installaties bepalen daarna welke elektrische of brandstofinput nodig is om die thermische vraag te leveren. "
        "Deze scheiding is essentieel om gasloosheid en netcapaciteit tegelijk te kunnen beoordelen."
    )

    doc.add_heading("3. Deelmodellen", level=1)

    add_model_section(
        doc,
        "3.1 Weerdata en simulatie-index",
        purpose=(
            "De simulatie draait op de tijdstappen uit de aangeleverde weerdata. Dit voorkomt dat de app ongemerkt "
            "kunstmatige of geinterpoleerde weersituaties gebruikt."
        ),
        code_basis=(
            "run_energy_system_simulation() bepaalt de index uit de weather DataFrame. Het gebouwmodel en PV-model eisen "
            "dezelfde DatetimeIndex; ontbrekende of dubbele timestamps worden niet in het model zelf gerepareerd."
        ),
        method=[
            "De index wordt timezone-aware gemaakt als dat nodig is, met standaard Europe/Amsterdam.",
            "Het gebouwmodel vereist t_amb_C en controleert optioneel wind_ms en ghi_Wm2.",
            "PV vereist ghi_Wm2 en gebruikt t_amb_C als die beschikbaar is voor temperatuurcorrectie.",
            "Als index of verplichte kolommen niet exact kloppen, stopt het model met een foutmelding in plaats van stilzwijgend te interpoleren.",
        ],
        formulas=[],
        outputs=["DatetimeIndex", "t_amb_C", "ghi_Wm2", "wind_ms"],
        motivation=(
            "Voor netpieken en PV-opbrengst is timing bepalend. Strikte indexcontrole is daarom belangrijker dan een schijnbaar vloeiende, "
            "maar mogelijk foutief aangevulde tijdreeks."
        ),
    )

    add_model_section(
        doc,
        "3.2 Gebouwmodel: thermische warmtevraag en koelvraag",
        purpose=(
            "Het gebouwmodel berekent hoeveel warmte of koeling het gebouw thermisch nodig heeft. Het model bepaalt dus de vraag, "
            "niet welke installatie die vraag levert."
        ),
        code_basis=(
            "simulate_thermal_demand() in gebouwmodel.py berekent Q_heat_kWth en Q_cool_kWth uit buitentemperatuur, setpoints, "
            "gebouwschil, ventilatie, infiltratie, interne warmte en zonnewinsten."
        ),
        method=[
            "Per tijdstap wordt bepaald of het gebouw in gebruik is op basis van dagen en gebruiksuren.",
            "Tijdens en buiten gebruik kunnen aparte verwarmings- en koelsetpoints gelden.",
            "Een deadband voorkomt tegelijk verwarmen en koelen rond de setpoints.",
            "Transmissieverlies wordt berekend met UA, opgebouwd uit U-waarden en oppervlakken van gevel, dak, vloer en glas.",
            "Ventilatie- en infiltratiestromen worden omgerekend naar voelbare warmtestroom met luchtdebiet, luchtdichtheid en soortelijke warmte.",
            "Interne warmtelast en zoninstraling via glas worden als winsten meegenomen. Bij verwarming verlagen ze de vraag; bij koeling verhogen ze de vraag.",
        ],
        formulas=[
            "UA = U_wall*A_wall + U_roof*A_roof + U_ground*A_ground + U_window*A_window",
            "Q_air = rho_air * cp_air * (Vdot / 3600) * DeltaT",
            "Q_solar = GHI * solar_proxy_factor * orientation_factor * A_window * g_value * shading_factor",
            "Q_heat = max(Q_trans + Q_vent + Q_inf - Q_internal - Q_solar, 0)",
            "Q_cool = max(Q_internal + Q_solar + Q_trans + Q_vent + Q_inf, 0)",
        ],
        outputs=["Q_heat_kWth", "Q_cool_kWth", "regime_heat", "regime_cool"],
        motivation=(
            "Dit is een transparant eerste-orde gebouwmodel. Het volgt de bouwfysische logica van warmtebalansen, maar vermijdt de complexiteit "
            "van een volledig dynamisch zone- of installatiemodel. Dat past bij een consultants-tool waarin invoer snel moet kunnen worden aangepast."
        ),
    )

    add_model_section(
        doc,
        "3.3 Elektrisch verbruik, processen en overig gebruik",
        purpose=(
            "Deze deelmodellen maken het reguliere elektrische gebruik zichtbaar los van warmte, opwek en opslag. Daardoor kan de gebruiker zien "
            "welke gebruikers bijdragen aan de basislast en pieken."
        ),
        code_basis=(
            "pelektro.py, pprocess.py en poverig.py gebruiken eenvoudige weekroosters. Elektrisch en overig gebruik werken met W/m2; processen "
            "werken met kW per proces of een standaard proces-/idle-vermogen."
        ),
        method=[
            "Zonder subloads gebruikt het model een bezet/niet-bezet vermogen over het gehele bruto vloeroppervlak.",
            "Met subloads wordt elk elektrisch of procesonderdeel apart berekend en daarna opgeteld.",
            "Processen hebben een actief vermogen en een idle-vermogen.",
            "Alle negatieve vermogens worden geweigerd, omdat passieve verbruikers geen opwek mogen worden.",
        ],
        formulas=[
            "P_electric(t) = p_occ_or_unocc(t) * BVO / 1000",
            "P_process(t) = sum_k P_process,k(t)",
            "P_base_electric = P_ref_heat + P_cooling + P_elektro + P_process + P_mobility + P_overig",
        ],
        outputs=["P_elektro_kW", "P_process_kW", "P_overig_kW", "P_electric_base_load_kW"],
        motivation=(
            "Een rooster-gebaseerd model is bewust gekozen omdat consultants vaak wel bedrijfsuren, oppervlak en grove vermogens kennen, "
            "maar niet altijd detailmetingen per apparaat hebben. Subloads maken verfijning mogelijk zodra meer informatie beschikbaar is."
        ),
    )

    add_model_section(
        doc,
        "3.4 Mobiliteit: elektrische voertuigen en slim laden",
        purpose=(
            "Het mobiliteitsmodel vertaalt het aantal aanwezige elektrische auto's naar een dagelijkse laadenergie en een laadprofiel. "
            "Daarmee wordt mobiliteit een expliciete elektriciteitsvraag in plaats van een constante kantoorurenlast."
        ),
        code_basis=(
            "pmobiliteit.py gebruikt aantal auto's, laadvermogen, accucapaciteit, aankomst-SoC, gewenste vertrek-SoC, aanwezigheid, "
            "laadvenster, locatiecap en laadmodus."
        ),
        method=[
            "De benodigde energie per auto is het verschil tussen gewenste vertrek-SoC en aankomst-SoC maal de gemiddelde accucapaciteit.",
            "De dagenergie is energie per auto maal het aantal aanwezige auto's.",
            "Direct laden vult de vraag vanaf aankomst zo snel mogelijk binnen laadvermogen en locatiecap.",
            "Slim laden bepaalt per tijdstap contractruimte als grid_cap - basislast en verdeelt de dagenergie over beschikbare ruimte.",
            "Als er onvoldoende contractruimte is, registreert het model niet-geladen mobiliteitsenergie.",
        ],
        formulas=[
            "E_car = C_battery * max(SoC_departure - SoC_arrival, 0) / 100",
            "E_day = E_car * n_cars * cars_present_fraction",
            "P_smart_allowed(t) = min(P_site_cap, n_present*P_charger, max(P_contract - P_base(t), 0))",
        ],
        outputs=["P_mobility_kW", "E_mobility_charged_kWh", "E_mobility_unserved_kWh"],
        motivation=(
            "De keuze voor een deterministisch energiebehoefte-model houdt de invoer begrijpelijk en voorkomt dat mobiliteit onrealistisch "
            "als constante last wordt behandeld. Slim laden sluit aan op het doel van de app: functioneren binnen een beperkte netaansluiting."
        ),
    )

    add_model_section(
        doc,
        "3.5 Zonnepanelen",
        purpose=(
            "Het PV-model berekent lokaal elektrisch vermogen uit globale zoninstraling, opgesteld vermogen, orientatie, helling, performance ratio, "
            "omvormerrendement en temperatuurcorrectie."
        ),
        code_basis=(
            "simulate_pv() in pv.py gebruikt ghi_Wm2, optioneel t_amb_C, installed_capacity_kWp, azimuth_deg, tilt_deg, performance_ratio, "
            "inverter_efficiency, temp_coeff_per_C en optioneel site_cap_kW."
        ),
        method=[
            "Globale horizontale instraling wordt genormaliseerd op 1000 W/m2 en begrensd.",
            "Orientatie wordt als factor tussen 0.6 en 1.0 toegepast, met zuid als referentie.",
            "Helling wordt als eenvoudige factor toegepast, met 35 graden als referentie.",
            "Celtemperatuur wordt benaderd als buitentemperatuur plus 0.03 maal GHI.",
            "Het vermogen wordt begrensd op een eventuele locatiecap en negatieve waarden worden verwijderd.",
        ],
        formulas=[
            "irr_factor = clip(GHI / 1000, 0, 1.5)",
            "orientation_factor = clip(1 - 0.4 * angular_delta_to_south / 180, 0.6, 1.0)",
            "tilt_factor = clip(1 - abs(tilt - 35) / 120, 0.75, 1.05)",
            "T_cell = T_amb + 0.03 * GHI",
            "P_pv = kWp * irr_factor * orientation_factor * tilt_factor * PR * inverter_efficiency * temp_factor",
        ],
        outputs=["P_pv_kW", "pv_irr_factor", "pv_temp_factor"],
        motivation=(
            "Dit is een vereenvoudigd PV-proxy model, geinspireerd door de PVWatts-benadering waarin instraling, temperatuur en systeemverliezen "
            "centraal staan. Het is gekozen omdat het weinig invoer vraagt en het effect van richting en seizoenen begrijpelijk maakt."
        ),
    )

    add_model_section(
        doc,
        "3.6 WKK: gecombineerde warmte en elektriciteit",
        purpose=(
            "De WKK levert tegelijk elektriciteit en warmte uit brandstof. Het model laat daardoor zien of WKK netimport kan verlagen, "
            "warmtevraag kan dekken of juist onbenutte warmte produceert."
        ),
        code_basis=(
            "dispatch_wkk() ondersteunt elektriciteitsgestuurd, warmtevraaggestuurd, hybride piekverlaging, warmtevraag met elektrische cap, "
            "must-run en uitgeschakeld."
        ),
        method=[
            "Het gekozen dispatchprofiel bepaalt het elektrische doelvermogen.",
            "Bij warmtevraaggestuurde dispatch wordt warmtevraag omgerekend naar het elektrische niveau dat nodig is om die warmte te leveren.",
            "Het vermogen wordt begrensd door nominaal WKK-vermogen en minimale deellast.",
            "Brandstofinput volgt uit elektrisch rendement; warmteproductie volgt uit thermisch rendement.",
            "Warmte wordt gebruikt tot de actuele warmtevraag; overschot wordt als gedumpte WKK-warmte geregistreerd.",
        ],
        formulas=[
            "P_required_for_heat = Q_heat_demand * eta_el / eta_th",
            "F_wkk = P_wkk_el / eta_el",
            "Q_wkk_available = F_wkk * eta_th",
            "Q_wkk_used = min(Q_wkk_available, Q_heat_demand)",
            "Q_wkk_dumped = max(Q_wkk_available - Q_wkk_used, 0)",
        ],
        outputs=["P_wkk_el_kW", "Q_wkk_used_kWth", "Q_wkk_dumped_kWth", "F_wkk_fuel_kW"],
        motivation=(
            "WKK is meegenomen omdat het in bestaande gebouwen vaak relevant is als overgangs- of piekoplossing. Door warmte en elektriciteit "
            "tegelijk te boeken wordt zichtbaar of een WKK energetisch nuttig is of vooral overschotten veroorzaakt."
        ),
    )

    add_model_section(
        doc,
        "3.7 Warmteketen: warmtepomp, ketel, warmtenet en referentie",
        purpose=(
            "De warmteketen vult resterende warmtevraag in nadat WKK-warmte en warmteopslag zijn toegepast. Hiermee wordt expliciet zichtbaar "
            "welke warmtebron de vraag dekt en hoeveel elektriciteit of brandstof daarvoor nodig is."
        ),
        code_basis=(
            "total.py stuurt de volgorde aan. dispatch_heat_pump(), dispatch_boiler() en dispatch_district_heat() leveren elk hun deel "
            "van de resterende warmtevraag. Een referentie-elektrische verwarming kan overblijvende vraag via COP invullen."
        ),
        method=[
            "De warmtepomp levert maximaal de ingestelde thermische capaciteit en gebruikt COP om elektriciteit te berekenen.",
            "COP kan vast, seizoensafhankelijk of weersafhankelijk zijn. De weersafhankelijke COP stijgt bij mildere buitentemperatuur en daalt bij kou.",
            "Als de warmtepomp gedeeld is met het contractvermogen, beperkt de beschikbare netruimte het elektrische warmtepompvermogen.",
            "De ketel dekt resterende warmtevraag binnen thermische capaciteit en rendement; gasinput wordt alleen geboekt bij brandstoftype gas.",
            "Warmtenet dekt resterende warmtevraag binnen capaciteit zonder lokale brandstofinput.",
            "Referentie-elektrische verwarming vult eventueel resterende warmtevraag in via seizoens-COP, zodat scenario's zonder expliciete warmtebron toch sluitend kunnen worden doorgerekend.",
        ],
        formulas=[
            "COP_weather = max(COP_nominal + 0.06 * (T_amb - 7), 1.0)",
            "P_hp_el = Q_hp_th / COP",
            "Q_hp_th = min(Q_remaining, Q_hp_capacity, COP * P_grid_headroom)",
            "F_boiler = Q_boiler_th / eta_boiler",
            "P_reference_heat_el = Q_unserved_final / COP_reference",
        ],
        outputs=["Q_hp_th_kWth", "P_hp_el_kW", "Q_boiler_th_kWth", "Q_dh_th_kWth", "P_heat_ref_el_kW"],
        motivation=(
            "De splitsing tussen gebouwvraag en installatie-invulling voorkomt dubbele COP/EER-logica. Dit is inhoudelijk belangrijk voor gasloosheidsanalyse: "
            "een gebouw heeft warmtevraag, maar de gekozen installatie bepaalt of die vraag elektrisch, via brandstof of via warmtenet wordt geleverd."
        ),
    )

    add_model_section(
        doc,
        "3.8 Warmteopslag",
        purpose=(
            "Warmteopslag vangt warmteoverschot op, vooral WKK-warmte die op dat moment niet direct kan worden gebruikt, en levert later warmte aan resterende vraag."
        ),
        code_basis=(
            "simulate_thermal_storage() gebruikt capaciteit, laad-/ontlaadvermogen, minimale/initiele/maximale SoC, laad-/ontlaadrendement en stilstandsverlies."
        ),
        method=[
            "Per tijdstap wordt eerst stilstandsverlies van de opgeslagen energie afgehaald.",
            "Daarna laadt de opslag uit thermisch overschot binnen vermogen, rendement en vrije capaciteit.",
            "Vervolgens ontlaadt de opslag naar resterende warmtevraag binnen vermogen, rendement en beschikbare energie boven minimale SoC.",
            "Overschot dat niet kan worden opgeslagen blijft als warmteoverschot zichtbaar; vraag die niet kan worden geleverd blijft ongedekt voor volgende warmtebronnen.",
        ],
        formulas=[
            "E_loss = E_storage * loss_factor_per_hour * dt",
            "Q_charge = min(Q_surplus, P_charge_max, free_capacity / (eta_charge * dt))",
            "Q_discharge = min(Q_deficit, P_discharge_max, available_energy * eta_discharge / dt)",
            "E_next = E - E_loss + Q_charge * eta_charge * dt - Q_discharge / eta_discharge * dt",
        ],
        outputs=["Q_thermal_storage_charge_kWth", "Q_thermal_storage_discharge_kWth", "E_thermal_storage_kWhth"],
        motivation=(
            "Een SoC-model is nodig omdat warmteopslag tijdsafhankelijk is: een overschot in de ochtend kan alleen later worden gebruikt als er capaciteit, "
            "rendement en voldoende resterende opslagenergie beschikbaar zijn."
        ),
    )

    add_model_section(
        doc,
        "3.9 Batterij",
        purpose=(
            "De batterij verschuift elektrische energie in de tijd. In de huidige code is de batterij geen economische optimizer, maar een regelgebaseerd "
            "zelfconsumptie- en piekreductiemodel."
        ),
        code_basis=(
            "simulate_battery() gebruikt P_residual_before_battery_kW, capaciteit, laad-/ontlaadvermogen, SoC-grenzen, roundtrip efficiency en laadstrategie."
        ),
        method=[
            "Het model splitst roundtrip efficiency symmetrisch over laden en ontladen via de wortel van het roundtriprendement.",
            "Bij positieve restvraag ontlaadt de batterij om netimport te verlagen, begrensd door ontlaadvermogen en beschikbare energie.",
            "Bij lokaal overschot kan de batterij laden volgens strategie 'alleen lokaal overschot'.",
            "Bij strategie 'laden tot contractruimte' mag de batterij ook uit het net laden zolang netto import onder contractvermogen blijft.",
            "SoC blijft tussen minimale en maximale fractie van de ingestelde capaciteit.",
        ],
        formulas=[
            "eta_charge = eta_discharge = sqrt(eta_roundtrip)",
            "P_discharge = min(P_residual, P_discharge_max, available_energy * eta_discharge / dt)",
            "P_charge_surplus = min(-P_residual, P_charge_max, room / (eta_charge * dt))",
            "P_charge_headroom = min(P_contract - P_residual, P_charge_max, room / (eta_charge * dt))",
        ],
        outputs=["P_battery_charge_kW", "P_battery_discharge_kW", "E_battery_kWh", "battery_soc_pct"],
        motivation=(
            "Regelgebaseerde batterijlogica is transparant en past bij het doel om snel te zien wat opslag doet voor netpieken en lokale benutting. "
            "Voor financiele optimalisatie of prijsarbitrage zou later een optimalisatiemodel nodig zijn."
        ),
    )

    add_model_section(
        doc,
        "3.10 Elektrische netbalans en stoplicht",
        purpose=(
            "De netbalans vertaalt alle vraag, opwek en batterijstromen naar netimport, teruglevering en contractoverschrijding. Dit is het centrale "
            "resultaat voor de vraag of het gebouw binnen de netaansluiting past."
        ),
        code_basis=(
            "total.py berekent de restbalans voor en na batterij. grid.py berekent pieken, jaarimport/export, percentielen, duurcurve, overschrijdingsuren "
            "en stoplichtstatus."
        ),
        method=[
            "Eerst wordt totale elektrische vraag bepaald inclusief warmtepomp en referentieverwarming.",
            "Daarna worden PV en WKK-elektriciteit afgetrokken.",
            "De batterij mag de restvraag verlagen of overschot opnemen.",
            "Positieve restbalans is netimport; negatieve restbalans is teruglevering.",
            "Het stoplicht vergelijkt piekimport, p99-import en overschrijdingsduur/-energie met het contractvermogen en ingestelde marges.",
        ],
        formulas=[
            "P_generation_total = P_pv + P_wkk_el",
            "P_residual_after_battery = P_residual_before_battery + P_battery_charge - P_battery_discharge",
            "contract_excess = max(P_grid_import - P_contract, 0)",
            "load_factor = annual_grid_import / (peak_grid_import * hours_year)  [conceptuele KPI]",
        ],
        outputs=["P_grid_import_kW", "P_grid_export_kW", "P_grid_contract_excess_kW", "grid_evaluation"],
        motivation=(
            "Piekvermogen, percentielen en duurcurves zijn praktischer voor netcapaciteit dan alleen jaarverbruik. Een korte overschrijding vraagt een andere "
            "maatregel dan een structureel tekort aan contractvermogen."
        ),
    )

    add_model_section(
        doc,
        "3.11 Modelchecks en meetvalidatie",
        purpose=(
            "Modelchecks bewaken of de simulatie fysisch en administratief logisch blijft. Meetvalidatie vergelijkt modeluitkomsten met werkelijke data "
            "als die beschikbaar zijn."
        ),
        code_basis=(
            "_compute_balance_checks() controleert warmtebalans, negatieve waarden en capaciteitsgrenzen. calibration.py lijnt meetdata en simulatie uit "
            "en berekent RMSE, MAE, MBE, NMBE, CV(RMSE), R2 en Pearson-correlatie."
        ),
        method=[
            "De warmtebalans vergelijkt warmtevraag met geleverde warmte, ongedekte warmte, opslagstromen en warmteoverschot.",
            "Niet-fysische negatieve waarden worden gecontroleerd voor onder andere warmtevraag, brandstofinput en netstromen.",
            "Capaciteitschecks vergelijken berekende vermogens met ingestelde grenzen van warmtepomp, WKK, ketel en warmtenet.",
            "Validatie kan worden uitgevoerd voor netimport, netexport, elektrische load, gas of warmte.",
            "Meetdata en simulatie kunnen naar uur-, dag- of maandniveau worden geaggregeerd.",
        ],
        formulas=[
            "error = simulated - measured",
            "RMSE = sqrt(mean(error^2))",
            "NMBE = 100 * mean(error) / mean(measured)",
            "CV(RMSE) = 100 * RMSE / mean(measured)",
        ],
        outputs=["sanity_checks", "grid_evaluation", "validation metrics"],
        motivation=(
            "De app is bedoeld voor besluitvorming. Daarom moet niet alleen het scenarioresultaat zichtbaar zijn, maar ook of de berekening intern klopt "
            "en hoe goed het model aansluit op eventuele meetdata."
        ),
    )

    doc.add_heading("4. Aannames en beperkingen", level=1)
    add_paragraph(
        doc,
        "De energieplanner is een technisch-energetisch scenariomodel. Dat betekent dat het model gekozen instellingen doorrekent "
        "en laat zien wat daarvan de gevolgen zijn voor energiestromen, netbelasting, warmtevoorziening en opslag. Het model is dus "
        "geen automatische optimizer, geen kostenmodel en geen definitief installatieontwerp. De onderstaande aannames en beperkingen "
        "zijn belangrijk bij het lezen van resultaten en bij het bespreken van scenario's met klanten."
    )

    doc.add_heading("4.1 Algemene modelaannames", level=2)
    add_three_col_table(
        doc,
        ("Aanname", "Wat dit betekent", "Gevolg voor interpretatie"),
        [
            (
                "Een gebouw of locatie",
                "De simulatie kijkt naar een individueel gebouw of gebouwlocatie met een eigen vraag, opwek, opslag en netaansluiting.",
                "Resultaten zijn niet automatisch geldig voor een gebied, wijknet of collectief energiesysteem.",
            ),
            (
                "Tijdreeksmodel",
                "Alle energiestromen worden per tijdstap berekend op basis van de aangeleverde weerdata-index.",
                "Pieken, timing en seizoenen zijn leidend; jaarvolumes alleen zijn onvoldoende om netcapaciteit te beoordelen.",
            ),
            (
                "Deterministische simulatie",
                "Bij gelijke invoer geeft het model dezelfde uitkomst. Dispatchregels zijn vooraf gekozen.",
                "De app vindt niet vanzelf de economisch of technisch optimale configuratie.",
            ),
            (
                "Gescheiden energiedragers",
                "Thermische vraag, elektrische vraag en brandstofinput worden apart bijgehouden.",
                "Gasloosheid en netbelasting kunnen tegelijk worden beoordeeld zonder COP- of brandstofstromen te vermengen.",
            ),
            (
                "Vereenvoudigde fysica",
                "Gebouw, PV, warmtepomp, WKK en opslag zijn transparante benaderingen, geen detailmodellen.",
                "De uitkomst is geschikt voor verkenning en scenariovergelijking, niet als definitieve engineering-berekening.",
            ),
        ],
    )

    doc.add_heading("4.2 Beperkingen per modelonderdeel", level=2)
    add_three_col_table(
        doc,
        ("Onderdeel", "Belangrijkste beperking", "Praktische consequentie"),
        [
            (
                "Gebouwmodel",
                "Het gebouw wordt niet als volledig dynamisch thermisch massa-model gesimuleerd. Er is geen ruimtelijke zonering, geen uurlijkse regeling per ruimte en geen vochtbalans.",
                "Gebruik resultaten als indicatie van warmtevraag/koelvraag. Voor definitief comfort- of installatieontwerp blijft detailmodellering nodig.",
            ),
            (
                "Weerdata",
                "De code accepteert geen stille interpolatie binnen het gebouw- en PV-model; ontbrekende of foutieve weerdata blokkeren de simulatie.",
                "Dat verhoogt betrouwbaarheid, maar betekent dat invoerdata vooraf goed moeten worden opgeschoond.",
            ),
            (
                "Elektrische lasten",
                "Basislasten, processen en overig gebruik zijn rooster- en vermogen-gebaseerd. Gedragsvariatie, stochastiek en kortdurende startsstromen ontbreken.",
                "Piekvermogens kunnen onderschat worden als processen of apparaten in werkelijkheid kort en zwaar schakelen.",
            ),
            (
                "Mobiliteit",
                "Het EV-model gebruikt gemiddelde accucapaciteit, gemiddelde aankomst-/vertrek-SoC en aanwezigheid. Er is geen individueel voertuiggedrag.",
                "Het model is geschikt voor laadenergie en contractruimte, maar niet voor operationele laadpleinoptimalisatie.",
            ),
            (
                "PV",
                "Het PV-model gebruikt een vereenvoudigde GHI-proxy met orientatie-, helling-, performance- en temperatuurfactoren. Er is geen schaduw-, horizon- of POA-transpositiemodel.",
                "PV-opbrengst is bruikbaar voor scenariovergelijking; bij investeringsbesluiten moet een gespecialiseerde PV-studie volgen.",
            ),
            (
                "WKK",
                "WKK-dispatch volgt regelkeuzes zoals elektriciteitsvraag, warmtevraag of hybride piekverlaging. Start/stopkosten, onderhoud en emissies worden niet geoptimaliseerd.",
                "Resultaten tonen energetische inzet en overschot, niet automatisch economische haalbaarheid.",
            ),
            (
                "Warmtepomp",
                "COP is vast, seizoensafhankelijk of lineair weersafhankelijk. Aanvoertemperatuur, bronregime, defrost, deellastcurves en hydrauliek zijn niet expliciet gemodelleerd.",
                "Elektriciteitsvraag van warmte is indicatief. Voor selectie van een echte warmtepomp is leverancier-/ontwerpdata nodig.",
            ),
            (
                "Ketel en warmtenet",
                "Ketel en warmtenet leveren binnen capaciteit en rendement/capaciteit; er is geen gedetailleerde regeling, temperatuurtraject of tariefoptimalisatie.",
                "Deze modules zijn vooral bedoeld om resterende warmtevraag en brandstof-/warmtenetafhankelijkheid zichtbaar te maken.",
            ),
            (
                "Opslag",
                "Batterij en warmteopslag volgen SoC-, vermogen- en rendementsregels, maar geen marktprijsoptimalisatie of voorspellende regeling.",
                "Opslagresultaten tonen technische flexibiliteit en piekreductie, niet automatisch de beste businesscase.",
            ),
            (
                "Netbeoordeling",
                "Het stoplicht beoordeelt de gebouwzijdige netimport ten opzichte van contractvermogen. Netkwaliteit, spanningsval, congestiegebied en aansluitvoorwaarden van de netbeheerder zitten niet in het model.",
                "Een groen resultaat betekent energetisch passend binnen de ingestelde grens, geen formele netbeheerder-goedkeuring.",
            ),
        ],
    )

    doc.add_heading("4.3 Datakwaliteit en onzekerheid", level=2)
    add_paragraph(
        doc,
        "De kwaliteit van de uitkomst wordt sterk bepaald door de kwaliteit van de invoer. Vooral gebruiksuren, oppervlaktes, U-waarden, ventilatie, "
        "procesvermogens, laadgedrag, contractvermogen en weerdata hebben grote invloed op pieken en jaarvolumes. Als deze waarden onzeker zijn, "
        "moet het resultaat worden gelezen als bandbreedte. In de praktijk is het verstandig om ten minste een conservatief scenario, een verwacht "
        "scenario en een ambitieus scenario naast elkaar te zetten."
    )
    add_bullets(
        doc,
        [
            "Gebruik meetdata waar beschikbaar om basislast, gasverbruik en pieken te controleren.",
            "Controleer of de zwaarste netweek en zwaarste warmteweek logisch passen bij het bedrijfstype.",
            "Behandel ontbrekende technische gegevens expliciet als aannames in klantgesprekken.",
            "Gebruik modelchecks niet als formaliteit: een warmtebalansfout of capaciteitsoverschrijding kan wijzen op verkeerde interpretatie van een scenario.",
        ],
    )

    doc.add_heading("5. Interpretatie van resultaten", level=1)
    add_paragraph(
        doc,
        "De resultatenpagina moet worden gelezen als een technisch verhaal in vaste volgorde. Eerst: past het scenario binnen de netaansluiting? "
        "Daarna: waardoor ontstaan de pieken? Vervolgens: is de warmtevoorziening dekkend en gasloos? Tot slot: welke maatregel draagt het meest bij "
        "aan lagere netbelasting, minder brandstofgebruik of betere benutting van lokale opwek?"
    )

    doc.add_heading("5.1 Beslissingssamenvatting en stoplicht", level=2)
    add_paragraph(
        doc,
        "Het stoplicht is een snelle beoordeling van de netbelasting ten opzichte van het ingestelde contractvermogen. Groen betekent dat piek en robuuste "
        "marge binnen de ingestelde grens blijven. Oranje betekent dat het scenario dicht op de grens zit of beperkt overschrijdt. Rood betekent dat de "
        "piekimport boven de toegestane grens komt en dat maatregelen of een andere aansluiting nodig zijn."
    )
    add_three_col_table(
        doc,
        ("KPI", "Wat het zegt", "Hoe te gebruiken"),
        [
            (
                "Piek netimport na batterij",
                "Hoogste elektrische import uit het net na lokale opwek en batterijdispatch.",
                "Vergelijk direct met contractvermogen. Dit is de belangrijkste netcapaciteits-KPI.",
            ),
            (
                "Piek boven contract",
                "Maximale overschrijding boven het ingestelde contractvermogen.",
                "Laat zien hoeveel piekreductie of extra contractruimte nodig is.",
            ),
            (
                "Uren boven contract",
                "Aantal uren waarin netimport boven contractvermogen ligt.",
                "Onderscheidt incidentele pieken van structureel tekort.",
            ),
            (
                "Jaarlijkse netimport",
                "Totale elektriciteit uit het net over het jaar.",
                "Relevant voor energiegebruik, maar minder bepalend voor aansluitcapaciteit dan de piek.",
            ),
            (
                "Ongedekte warmte",
                "Warmtevraag die door geen warmtebron of referentievoorziening is geleverd.",
                "Moet voor een technisch sluitend verwarmingsscenario normaal nul of verklaarbaar zijn.",
            ),
            (
                "Gas-/brandstofinput",
                "Brandstof die door WKK en ketel wordt gebruikt.",
                "Bepaalt of een scenario werkelijk gasloos of brandstofarm is.",
            ),
        ],
    )

    doc.add_heading("5.2 Netcapaciteit lezen", level=2)
    add_paragraph(
        doc,
        "Netcapaciteit moet altijd met tijdreeksen worden gelezen. Een jaarvolume kan laag zijn terwijl een korte piek toch het contractvermogen overschrijdt. "
        "De zwaarste netweek toont wanneer de hoogste netbelasting ontstaat en welke gebruikers of installaties eraan bijdragen. De duurcurve laat zien "
        "hoe vaak hoge waarden voorkomen."
    )
    add_bullets(
        doc,
        [
            "Een enkele korte piek wijst vaak op regelstrategie, slim laden, batterij of processturing.",
            "Veel uren boven contract wijzen eerder op te weinig aansluitvermogen, te veel gelijktijdige elektrificatie of onvoldoende lokale opwek/flexibiliteit.",
            "Een hoge p99-waarde betekent dat de belasting niet alleen door een uitschieter wordt veroorzaakt.",
            "Gemiddelde netruimte is nuttig voor slim laden of batterij laden, maar zegt niet vanzelf dat elke piek oplosbaar is.",
        ],
    )

    doc.add_heading("5.3 Verbruik, opwek en load match", level=2)
    add_paragraph(
        doc,
        "De jaar- en maandprofielen laten zien of vraag en opwek op dezelfde momenten plaatsvinden. Dit is vooral belangrijk bij PV: veel opwek in de zomer "
        "helpt beperkt bij winterse warmtepomppieken. Load match gaat daarom niet alleen over hoeveel PV wordt opgewekt, maar ook over wanneer die opwek "
        "beschikbaar is ten opzichte van gebouwvraag, mobiliteit en processen."
    )
    add_three_col_table(
        doc,
        ("Resultaat", "Interpretatie", "Typische maatregel"),
        [
            (
                "Veel teruglevering midden op de dag",
                "Lokale opwek valt niet samen met lokale vraag of opslag is te klein/vol.",
                "Batterij, slim laden, procesverschuiving of andere PV-orientatie onderzoeken.",
            ),
            (
                "Hoge ochtend- of avondpiek",
                "Vraag ligt buiten PV-productieuren of laadt/processen starten tegelijk.",
                "Slim laden, starttijden spreiden, batterij ontladen op piekmomenten.",
            ),
            (
                "Hoge wintervraag",
                "Warmtepomp of referentieverwarming kan netpiek veroorzaken terwijl PV laag is.",
                "Warmteopslag, grotere warmtebron, lagere warmtevraag of hybride bron onderzoeken.",
            ),
            (
                "Laag eigengebruik van PV",
                "PV-opwek wordt relatief vaak geexporteerd.",
                "Opslag of vraagsturing kan waarde toevoegen; extra PV helpt niet altijd tegen netpiek.",
            ),
        ],
    )

    doc.add_heading("5.4 Warmte en gasloosheid lezen", level=2)
    add_paragraph(
        doc,
        "Een gasloos scenario is pas overtuigend als de warmtevraag door niet-gasbronnen wordt gedekt en de elektrische consequentie daarvan binnen de "
        "netaansluiting past. Alleen brandstofinput verlagen is dus niet genoeg; de warmtepomp kan de gasvraag vervangen door elektrische piekvraag."
    )
    add_bullets(
        doc,
        [
            "Warmtevraag is de behoefte van het gebouw; warmtelevering is wat installaties daadwerkelijk invullen.",
            "Warmtepomp-elektriciteit moet worden meegelezen in de netcapaciteit, vooral in koude weken.",
            "Ongedekte warmte betekent dat het scenario technisch niet volledig voorziet in de warmtevraag, tenzij dit bewust is toegestaan.",
            "WKK-warmteoverschot is niet hetzelfde als ongedekte warmte: overschot betekent dat warmte beschikbaar was maar op dat moment niet nuttig kon worden gebruikt.",
            "Gasinput uit WKK of ketel betekent dat het scenario niet volledig gasloos is, ook als de netbelasting gunstig is.",
        ],
    )

    doc.add_heading("5.5 Opslag en flexibiliteit lezen", level=2)
    add_paragraph(
        doc,
        "Opslag moet worden beoordeeld op timing en benutting. Een batterij of warmtebuffer is nuttig wanneer die laadt op momenten met overschot of netruimte "
        "en ontlaadt op momenten met piek of warmtetekort. Alleen een grote capaciteit is geen garantie voor effect als laad- en ontlaadmomenten niet passen."
    )
    add_bullets(
        doc,
        [
            "Als de batterij vaak leeg is tijdens pieken, is capaciteit of laadmoment onvoldoende.",
            "Als de batterij vaak vol is terwijl er nog teruglevering optreedt, is vermogen/capaciteit te klein of vraagsturing nodig.",
            "Batterij laden is altijd hoger of gelijk aan nuttige ontlading gedeeld door rendement; verschil tussen laden en ontladen is normaal door verliezen en SoC-eindstand.",
            "Warmteopslag helpt vooral wanneer warmteoverschot en warmtetekort binnen een passende tijdsafstand liggen.",
        ],
    )

    doc.add_heading("5.6 Modelchecks lezen", level=2)
    add_paragraph(
        doc,
        "Modelchecks zijn bedoeld als kwaliteitsfilter. Een waarschuwing betekent niet automatisch dat alle resultaten onbruikbaar zijn, maar wel dat de gebruiker "
        "moet begrijpen wat er aan de hand is voordat conclusies worden getrokken."
    )
    add_three_col_table(
        doc,
        ("Check", "Betekenis", "Actie bij aandacht nodig"),
        [
            (
                "Warmtebalans",
                "Controleert of warmtevraag, geleverde warmte, opslag, overschot en ongedekte warmte administratief sluiten.",
                "Controleer warmtebronnen, opslagstromen en of WKK-overschot verkeerd als residu wordt gelezen.",
            ),
            (
                "Negatieve fysica",
                "Controleert of niet-negatieve grootheden zoals brandstofinput, netimport en warmtevraag niet onder nul komen.",
                "Controleer invoerwaarden, projectbestand en eventuele ontbrekende kolommen.",
            ),
            (
                "Vermogensgrenzen",
                "Controleert of installaties boven ingestelde capaciteit leveren.",
                "Controleer capaciteit, deellastinstellingen en dispatchvolgorde.",
            ),
            (
                "Validatie met meetdata",
                "Vergelijkt model en meting met foutmaten zoals RMSE, NMBE en CV(RMSE).",
                "Kalibreer basislasten, gebruiksuren, processen of warmte-instellingen voordat scenario's worden beoordeeld.",
            ),
        ],
    )

    doc.add_heading("6. Bronnen en onderbouwing", level=1)
    add_paragraph(
        doc,
        "De onderstaande bronnen onderbouwen de gekozen modelindicatoren en de vereenvoudigde rekenaanpak. Ze worden in een volgende iteratie "
        "per deelmodel nog preciezer gekoppeld aan de definitieve tekst en app-uitleg."
    )
    add_bullets(
        doc,
        [
            "IEA, Energy End-uses and Efficiency Indicators: onderbouwt het gebruik van eindgebruikcategorieen en energie-intensiteit als kernindicatoren voor gebouwen. https://www.iea.org/data-and-statistics/data-tools/energy-end-uses-and-efficiency-indicators-data-explorer",
            "IEA, Buildings en The Future of Heat Pumps: onderbouwt het belang van elektrificatie, warmte/koude als eindgebruik en COP als verhouding tussen geleverde warmte en elektrische input. https://www.iea.org/reports/buildings en https://www.iea.org/reports/the-future-of-heat-pumps/how-a-heat-pump-works",
            "EnergyPlus Engineering Reference: bouwfysische achtergrond voor warmtebalansen, interne winsten, ventilatie, infiltratie en zonnewinsten. https://energyplus.net/documentation",
            "NREL/PVPMC PVWatts: onderbouwt een PV-model waarin instraling, systeemverliezen en temperatuurcorrectie bepalend zijn voor PV-productie. https://pvpmc.sandia.gov/modeling-guide/2-dc-module-iv/point-value-models/pvwatts/",
            "PV-load matching literatuur: onderbouwt indicatoren zoals zelfconsumptie, zelfvoorziening, timing tussen opwek en vraag en de rol van opslag. https://www.sciencedirect.com/science/article/pii/S0306261919309110",
            "NREL REopt en opslagpublicaties: onderbouwen het gebruik van opslag voor piekreductie, zelfconsumptie en dispatchbeslissingen. https://www.nrel.gov/reopt/",
            "ASHRAE Guideline 14 / FEMP-calibratiecriteria: onderbouwt het gebruik van NMBE en CV(RMSE) voor vergelijking tussen model en meetdata. https://pmc.ncbi.nlm.nih.gov/articles/PMC7506729/",
        ],
    )

    doc.add_heading("Open punten voor volgende iteratie", level=1)
    add_bullets(
        doc,
        [
            "De tekst per deelmodel inhoudelijk reviewen en waar nodig inkorten of uitbreiden voor de uiteindelijke app-weergave.",
            "Bronnen per paragraaf preciezer verwerken met verwijzingen in lopende tekst.",
            "De modelcheck voor onbenutte WKK-warmte scherper formuleren, zodat overtollige warmte niet wordt verward met ongedekte warmte.",
            "Bepalen welke tekst uiteindelijk in de app zelf komt en welke tekst alleen in het downloadbare methodedocument blijft.",
        ],
    )

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
